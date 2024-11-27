import streamlit as st
import geopandas as gpd
import pandas as pd
import plotly.express as px
import matplotlib.pyplot as plt
import json
from groq import Groq
from shapely.geometry import MultiPolygon, Polygon
import numpy as np
import plotly.graph_objects as go
import pandas as pd
import nltk
from nltk.tokenize import word_tokenize
from groq import Groq
import streamlit as st
import pandas as pd
import geopandas as gpd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import base64
import plotly.express as px
from plotly.subplots import make_subplots

def generar_informes(selected_municipio):
    #importamos los archivos
    mapa_base = gpd.read_file('DatosU/DatosU/mapa_quindio/mapa_quindio.shp')
    evas_quindio = pd.read_excel('DatosU/DatosU/EVAS_Quindio.xlsx')
    ipm_quindio = pd.read_excel('DatosU/DatosU/IPM_Quindio.xlsx')
    general_quindio = pd.read_excel('DatosU/DatosU/datos_quindio_generales.xlsx')
    socioeconomico_quindio = pd.read_excel('DatosU/DatosU/Socioeconomico_Quindio.xlsx')

    #Seleccionamos el municipio
    municipio_seleccionado = str(selected_municipio)
    #Ubicación de los municipios en el Departamento del Quindío
    ubicaciones = {
        'Armenia': 'centro',
        'Buenavista': 'suroccidente',
        'Calarcá': 'oriente',
        'Circasia': 'centro-norte',
        'Córdoba': 'suroriente',
        'Filandia': 'norte',
        'Génova': 'sur',
        'La Tebaida': 'occidente',
        'Montenegro': 'occidente',
        'Pijao': 'sur',
        'Quimbaya': 'noroccidente',
        'Salento': 'nororiente',
        }

    #### INTRODUCCIÓN ####


    #Definir las variables de la introducción
    #Extensión
    datos_generales_quindio = general_quindio[general_quindio['Municipio'] == municipio_seleccionado]
    extension = datos_generales_quindio['Área en Km2'].values[0]

    #Población total 2022

    dato_población_quindio = socioeconomico_quindio[(socioeconomico_quindio['Municipio'] == municipio_seleccionado)
                                                        & (socioeconomico_quindio['Año'] == 2022)]

    poblacion_total = dato_población_quindio['Población Total'].values[0]

    #Crear la figura y los ejes
    plt.figure(figsize=(12,10))
    ax = plt.axes()
    mapa_base.plot(ax=ax, color='white', edgecolor='#363636')

    #Filtrar el municipio de la variable de entrada
    nombre_municipio = mapa_base[mapa_base['MPIO_CNMBR'] == municipio_seleccionado]
    print(nombre_municipio)
    #Agregar el municipio en el mapa
    nombre_municipio.plot(ax=ax, color='green', edgecolor='black')

    #Agregar los detalles estéticos
    plt.title(f'Municipio de {municipio_seleccionado}', fontsize=20, fontweight='bold')
    ax.set_xlabel('Longitud', color='#363636', fontsize=14)
    ax.set_ylabel('Latitud', color='#363636', fontsize=14)
    plt.grid(color='grey', linestyle='--', linewidth=0.5)

    # Cambiar el color de todos los bordes (spines) a gris
    for spine in ax.spines.values():
        spine.set_color('#363636')

    # Configurar los parámetros de las marcas y la cuadrícula
    ax.tick_params(colors='0.2', grid_color='#515151')

    # Ajustar los márgenes de la figura
    plt.subplots_adjust(left=0.05, right=0.95, top=0.95, bottom=0.05)

    #Ajustar el tamaño de la figura
    plt.tight_layout()

    # Guardar la imagen
    plt.savefig('mapa.png')
    plt.close()

    #### PERFIL DEMOGRÁFICO ####

    # Aplicar el filtro del municipio seleccionado y del año
    datos_municipio_población = socioeconomico_quindio[(socioeconomico_quindio['Municipio'] == municipio_seleccionado)
                                                        & (socioeconomico_quindio['Año'] >= 2013)]

    # definir las variables
    poblacion_rural = dato_población_quindio['Población Rural'].values[0]
    porcentaje_poblacion_rural = round(poblacion_rural / poblacion_total *100, 2)

    poblacion_urbana = dato_población_quindio['Población Urbana'].values[0]
    porcentaje_poblacion_urbana = round(poblacion_urbana / poblacion_total *100, 2)

    #Gráfico


    #Crear la figura y los ejes
    plt.figure(figsize=(10, 5))
    ax = plt.axes()

    #Crear la gráfica de líneas para la población total
    ax.plot(datos_municipio_población['Año'], 
            datos_municipio_población['Población Total'], 
            color='royalblue', 
            marker='o', 
            linestyle='-', 
            linewidth=2, 
            markersize=8,
            label='Población Total')

    #Ajustar el ancho de las barras
    bar_width = 0.48

    #Calcular la posición de las barras
    posicion_rural = datos_municipio_población['Año'] - bar_width / 2
    posicion_urbana = datos_municipio_población['Año'] + bar_width / 2

    #Crear las barras para las poblaciones rural y urbana
    barras_rural = ax.bar(posicion_rural, 
                        datos_municipio_población['Población Rural'], 
                        bar_width, 
                        color='green', 
                        label='Población Rural')
    barras_urbana = ax.bar(posicion_urbana, 
                        datos_municipio_población['Población Urbana'], 
                        bar_width, 
                        color='grey', 
                        label='Población Urbana')

    # Agregar etiquetas a las barras
    for barra in barras_rural:
        altura = barra.get_height()
        ax.text(barra.get_x() + barra.get_width() / 2, altura - 0.09 * altura, f'{(altura/1000).round(1)}', ha='center', va='bottom', fontsize=9.5, color='black')
    for barra in barras_urbana:
        altura = barra.get_height()
        ax.text(barra.get_x() + barra.get_width() / 2, altura - 0.09 * altura, f'{(altura/1000).round(1)}', ha='center', va='bottom', fontsize=9.5, color='black')

    # Agregar etiquetas a la línea
    for x, y in zip(datos_municipio_población['Año'], datos_municipio_población['Población Total']):
        ax.annotate(f'{round(y/1000, 1)}', xy=(x, y), xytext=(0, 5), textcoords='offset points', ha='center', fontsize=10, color='royalblue')


    #Agregar los detalles estéticos
    plt.title(f'Evolución poblacional de {municipio_seleccionado}', fontsize=20, fontweight='bold', pad=40)

    #Configurar el eje x
    ax.set_xticks(datos_municipio_población['Año'])
    ax.set_xlabel('Año', color='black', fontsize=12)
    ax.tick_params(axis='x', length=0)
    ax.set_xticklabels(datos_municipio_población['Año'], fontsize=11)

    #Ocultar los valores y ticks del eje y
    ax.set_yticklabels([])
    ax.yaxis.set_ticks([])


    # Eliminar los bordes
    for spine in ax.spines.values():
        spine.set_visible(False)

    #Agregar y configurar la leyenda
    ax.legend(loc='upper center', fontsize=11, bbox_to_anchor=(0.5, 1.15), ncols=3)

    #Añadir texto con aclaración sobre los valores de población
    plt.text(0.21, 0.99, '*Población en miles de habitantes', transform=ax.transAxes, fontsize=9, color='#4B4B4B')

    #Mostrar el gráfico
    plt.tight_layout()

    # Guardar la imagen
    plt.savefig('grafico1.png')
    plt.close()


    #### PERFIL ECONÓMICO-PRODUCTIVO ####

    #### CULTIVOS

    #Filtrar los datos del municipio seleccionado y del año 2022
    datos_municipio_cultivos = evas_quindio[(evas_quindio['Municipio'] == municipio_seleccionado)
                                            & (evas_quindio['Año'] == 2022)]

    # Definir las variables
    # Obtener la fila con el valor máximo de 'Área sembrada (ha)'
    fila_cultivo_principal = datos_municipio_cultivos.loc[datos_municipio_cultivos['Área sembrada (ha)'].idxmax()]

    # Extraer el valor de hectáreas y el nombre del cultivo
    hectareas_cultivo_principal = fila_cultivo_principal['Área sembrada (ha)']
    cultivo_principal = fila_cultivo_principal['Cultivo']

    # Lista de cultivos
    lista_de_cultivos = datos_municipio_cultivos['Cultivo'].unique()
    lista_de_cultivos_str = ', '.join(lista_de_cultivos)

    # Extensión de cultivos
    suma_de_cultivos = datos_municipio_cultivos['Área sembrada (ha)'].sum()

    #Seleccionar los 10 cultivos con mayor área sembrada
    datos_municipio_cultivos = datos_municipio_cultivos.sort_values('Área sembrada (ha)', ascending=False).head(10)

    # Gráfico

    #Crear la figura y los ejes
    plt.figure(figsize=(10, 5))
    ax = plt.axes()

    #Crear las barras horizontales
    barras = ax.barh(datos_municipio_cultivos['Cultivo'],datos_municipio_cultivos['Área sembrada (ha)'], color='green', label='Área sembrada')

    #Organizar de mayor a menor los cultivos de acuerdo con su área sembrada
    plt.gca().invert_yaxis()

    # Agregar etiquetas a las barras
    for barra in barras:
        ancho = barra.get_width()
        ax.text(ancho, barra.get_y() + barra.get_height() / 2, f'{ancho:.1f}', ha='left', va='center', fontsize=10, color='black')

    #Agregar los detalles estéticos
    plt.title(f'Área sembrada (ha) por cultivo en {municipio_seleccionado}', fontsize=20, fontweight='bold')

    #Ocultar los valores y ticks del eje y
    ax.set_xticklabels([])
    ax.xaxis.set_ticks([])
    ax.tick_params(axis='y', length=0)

    # Eliminar los bordes
    for spine in ax.spines.values():
        spine.set_visible(False)

    #Agregar y configurar la leyenda
    plt.tight_layout()

    # Guardar la imagen
    plt.savefig('grafico2.png')
    plt.close()


    #### PIB MUNICIPAL

    # Aplicar el filtro del municipio seleccionado y del año
    datos_municipio_valor_agregado = socioeconomico_quindio[(socioeconomico_quindio['Municipio'] == municipio_seleccionado)
                                                            & (socioeconomico_quindio['Año'] >= 2013)]

    # Definir las variables
    datos_municipio_pib = socioeconomico_quindio[(socioeconomico_quindio['Municipio'] == municipio_seleccionado)
                                                            & (socioeconomico_quindio['Año'] == 2022)]

    #Extraer el valor del PIB
    PIB_primario = datos_municipio_pib['PIB Primario'].values[0]
    PIB_secundario = datos_municipio_pib['PIB Secundario'].values[0]
    PIB_terciario = datos_municipio_pib['PIB Terciario'].values[0]
    PIB_total = datos_municipio_pib['PIB'].values[0]

    #Porcentajes
    porcentaje_PIB_primario = round(PIB_primario / PIB_total * 100, 2)
    porcentaje_PIB_secundario = round(PIB_secundario / PIB_total * 100, 2)
    porcentaje_PIB_terciario = round(PIB_terciario / PIB_total * 100, 2)

    #Gráifico

    # Crear la figura y los ejes
    plt.figure(figsize=(10, 5))
    ax = plt.axes()

    # Crear las barras apiladas
    barras_primario = ax.bar(datos_municipio_valor_agregado['Año'], 
                            datos_municipio_valor_agregado['PIB Primario'], 
                            color='green', 
                            label='PIB Primario')

    barras_secundario = ax.bar(datos_municipio_valor_agregado['Año'], 
                            datos_municipio_valor_agregado['PIB Secundario'], 
                            bottom=datos_municipio_valor_agregado['PIB Primario'], 
                            color='grey', 
                            label='PIB Secundario')

    barras_terciario = ax.bar(datos_municipio_valor_agregado['Año'], 
                            datos_municipio_valor_agregado['PIB Terciario'], 
                            bottom=datos_municipio_valor_agregado['PIB Primario'] + datos_municipio_valor_agregado['PIB Secundario'], 
                            color='royalblue', 
                            label='PIB Terciario')

    # Etiquetas con porcentaje para cada barra
    for barras, columna in zip([barras_primario, barras_secundario, barras_terciario], 
                            ['PIB Primario', 'PIB Secundario', 'PIB Terciario']):
        for barra in barras:
            altura = barra.get_height()
            # Usamos el valor central de cada barra para encontrar el año
            año = barra.get_x() + barra.get_width() / 2
            # Obtener el índice correspondiente al año
            año_idx = datos_municipio_valor_agregado.index[(datos_municipio_valor_agregado['Año'] == año)].tolist()[0]
            total = datos_municipio_valor_agregado.loc[año_idx, 'PIB']
            porcentaje = (altura / total) * 100
            # Ajustar la posición vertical de la etiqueta
            pos_y = barra.get_y() + barra.get_height() / 2
            ax.text(barra.get_x() + barra.get_width() / 2, 
                    pos_y, 
                    f'{porcentaje:.1f}%', 
                    ha='center', 
                    va='center', 
                    fontsize=10, 
                    color='white')
        
    # Agregar los detalles estéticos
    plt.title(f'PIB por sector en {municipio_seleccionado}', fontsize=20, fontweight='bold', pad=40)

    # Configurar el eje x
    ax.set_xticks(datos_municipio_valor_agregado['Año'])
    ax.set_xlabel('Año', color='black', fontsize=12)
    ax.tick_params(axis='x', length=0)
    ax.set_xticklabels(datos_municipio_valor_agregado['Año'], fontsize=11)

    # Ocultar los valores y ticks del eje y
    ax.set_yticklabels([])
    ax.yaxis.set_ticks([])

    # Eliminar los bordes
    for spine in ax.spines.values():
        spine.set_visible(False)

    # Agregar y configurar la leyenda
    ax.legend(loc='upper center', fontsize=11, bbox_to_anchor=(0.5, 1.15), ncol=3)

    # Añadir texto con aclaración sobre los valores de población
    plt.text(0.24, 0.99, '*Porcentaje PIB Municipal', transform=ax.transAxes, fontsize=9, color='#4B4B4B')

    # Mostrar el gráfico
    plt.tight_layout()

    # Guardar la imagen
    plt.savefig('grafico3.png')
    plt.close()

    #### INGRESOS PREDIALES

    # Aplicar el filtro del municipio seleccionado y del año
    datos_municipio_ingresos_predial = socioeconomico_quindio[(socioeconomico_quindio['Municipio'] == municipio_seleccionado)
                                            & (socioeconomico_quindio['Año'] == 2022)]

    # definir las variables
    ingresos_predial = datos_municipio_ingresos_predial['Ingresos Prediales'].values[0]
    ingreso_total = datos_municipio_ingresos_predial['Ingresos Totales'].values[0]
    porcentaje_ingreso_predial = round(ingresos_predial / ingreso_total * 100, 2)

    #Gráfico

    # Crear la figura y los ejes
    plt.figure(figsize=(10, 5))
    ax = plt.axes()

    # Calcular los valores
    ingresos_totales = datos_municipio_ingresos_predial['Ingresos Totales'].sum()
    ingresos_prediales = datos_municipio_ingresos_predial['Ingresos Prediales'].sum()
    otros_ingresos = ingresos_totales - ingresos_prediales

    # Crear el gráfico de pastel
    ax.pie([ingresos_prediales, otros_ingresos], labels=['Ingresos Prediales', 'Otros Ingresos'], autopct='%1.1f%%', colors=['#00CC66','#0066FF'], startangle=90)
    ax.axis('equal')

    # Agregar los detalles estéticos
    plt.title(f'Proporción de Ingresos Prediales en {municipio_seleccionado}', fontsize=20, fontweight='bold')

    # Mostrar el gráfico
    plt.tight_layout()

    # Guardar la imagen
    plt.savefig('grafico4.png')
    plt.close()

    #### POBREZA MULTIDIMENSIONAL ####

    #Aplicar el fitlro del municipio seleccionado
    datos_municipio_pobreza = ipm_quindio[(ipm_quindio['Municipio'] == municipio_seleccionado)]

    #Definir las variables IPM general
    ipm_municipal = datos_municipio_pobreza['IPM'].values[0]
    IPM_Rural = datos_municipio_pobreza['IPM Rural'].values[0]
    IPM_Urbano = datos_municipio_pobreza['IPM Urbano'].values[0]

    #Gráfico

    # Crear la figura y los ejes
    plt.figure(figsize=(10, 5))
    ax = plt.axes()

    # Preparar los datos para el gráfico
    categorias = ['IPM Municipal', 'IPM Urbano', 'IPM Rural']
    valores = [datos_municipio_pobreza['IPM'].values[0], datos_municipio_pobreza['IPM Urbano'].values[0], datos_municipio_pobreza['IPM Rural'].values[0]]

    # Crear las barras horizontales
    ax.barh(categorias, valores, color=['#0066FF','#808080','#00CC99'])

    # Agregar el título
    plt.title(f'Índice de Pobreza Multidimensional en {municipio_seleccionado}', fontsize=20, fontweight='bold')

    # Agregar etiquetas a las barras
    for i, v in enumerate(valores):
        ax.text(v + 0.01, i, f'{v:.2f}', color='black', va='center', fontsize=10)

    #Ocultar los valores y ticks del eje y
    ax.set_xticklabels([])
    ax.xaxis.set_ticks([])
    ax.tick_params(axis='y', length=0)

    # Eliminar los bordes
    for spine in ax.spines.values():
        spine.set_visible(False)

    # Mostrar el gráfico
    plt.tight_layout()

    # Guardar la imagen
    plt.savefig('grafico5.png')
    plt.close()

    #Definir las variables mayores IPM
    trabajo_informal_municipio = datos_municipio_pobreza['Trabajo Informal'].values[0]
    trabajo_informal_rural = datos_municipio_pobreza['Trabajo Informal Rural'].values[0]
    trabajo_informal_urbano = datos_municipio_pobreza['Trabajo Informal Urbano'].values[0]


    ##### DOCUMENTO DE WORD ####

    doc = Document()

    # Ajustar los márgenes del documento
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Función para agregar párrafos justificados con tamaño de letra 12
    def add_paragraph(doc, text):
        paragraph = doc.add_paragraph(text)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for run in paragraph.runs:
            run.font.size = Pt(12)

    # Agregar contenido al documento
    doc.add_heading(f'Caracterización Territorial del Municipio {municipio_seleccionado}, Quindío', 0)
    doc.add_heading('Introducción', level=1)
    add_paragraph(doc, f'El municipio de {municipio_seleccionado} se encuentra ubica en el {ubicaciones[municipio_seleccionado]} del Departamento tiene una extensión de {extension} km2 y una población de {poblacion_total} en el año 2022 (DANE,2022). A continuación, se muestra su ubicación en el Departamento:')
    doc.add_picture('mapa.png', width=Inches(6.5))
    doc.add_heading('Perfil Demográfico', level=1)
    add_paragraph(doc, f'De acuerdo con el DANE, {municipio_seleccionado} al año 2022 tiene una población de {poblacion_total} que se distribuye en {poblacion_rural} habitantes rurales, equivalente al {porcentaje_poblacion_rural}% del total y una población urbana de {poblacion_urbana} habitantes el {porcentaje_poblacion_urbana}%. La gráfica 1 muestra la evolución de la población de {municipio_seleccionado} entre los años 2013 a 2022:')
    doc.add_picture('grafico1.png', width=Inches(6.5))
    doc.add_heading('Perfil Económico-Productivo', level=1)
    doc.add_heading('Cultivos', level=2)
    add_paragraph(doc, 'La economía del departamento del Quindío se caracteriza por su orientación a los sectores primario y terciario de la economía, específicamente al desarrollo de actividades agropecuarias y la prestación de servicios enfocados en el turismo y sus derivados.')
    add_paragraph(doc, f'De acuerdo con las Evaluaciones Agropecuarias Municipales, EVAS, del Ministerio de Agricultura y Desarrollo Rural y la UPRA, los cultivos presentes en {municipio_seleccionado}, durante el año 2022, son {lista_de_cultivos_str} con una extensión conjunta de {suma_de_cultivos} hectáreas, de estos el principal es {cultivo_principal} con una extensión de {hectareas_cultivo_principal} hectáreas. La gráfica 2 presenta principales cultivos al año 2022, de acuerdo con su área sembrada:')
    doc.add_picture('grafico2.png', width=Inches(6.5))
    doc.add_heading('PIB Municipal', level=2)
    add_paragraph(doc, f'Las Cuentas Nacionales del DANE estiman el Producto Interno Bruto Municipal a partir del valor agregado generado por las actividades en los tres sectores de la economía. En el caso de {municipio_seleccionado}, durante el año 2022, el sector primario representó el {porcentaje_PIB_primario}%, el sector secundario tuvo un peso del {porcentaje_PIB_secundario}% y el sector terciario generó el {porcentaje_PIB_terciario}% del total. La gráfica 3 muestra la evolución del PIB entre los años 2013 a 2022, con el porcentaje de cada uno de los tres sectores:')
    doc.add_picture('grafico3.png', width=Inches(6.5))
    doc.add_heading('Ingresos Prediales', level=2)
    add_paragraph(doc, f'Los Ingresos Municipales dependen de diversos factores, transferencias nacionales, actividades económicas del municipio, regalías, entre otros, sin embargo, para la mayoría de los municipios de sexta categoría (la predominante en Quindío), el ingreso por impuesto predial representa la mayor parte de sus Ingresos Corrientes de Libre Destinación, es decir, aquellos de los que pueden disponer para realizar las inversiones que requiere su municipio y de los cuales tiene libertad para decidir su destinación. Para el caso {municipio_seleccionado}, los ingresos por predial representan el {porcentaje_ingreso_predial}% de los ingresos totales como se observa en el gráfico 4:')
    doc.add_picture('grafico4.png', width=Inches(6.5))
    doc.add_heading('Pobreza Multidimensional', level=2)
    add_paragraph(doc, 'El DANE mide el Indicador de Pobreza Multidimensional, IPM, a partir de las privaciones que tienen los hogares. A partir de los datos del Censo 2018, el DANE, realizó mediciones precisas de las diferentes dimensiones que componen este indicador brindando un panorama muy detallado de las atenciones que necesitan los municipios. El indicador se mide de 1 a 100, representando el porcentaje de las personas que se encuentran en condición de privación.')
    add_paragraph(doc, f'{municipio_seleccionado} tiene un IPM de {ipm_municipal}, sin embargo, al territorializarlo se identifica que la situación es más difícil en el sector rural con un IPM de {IPM_Rural} en comparación con el IPM urbano que es de {IPM_Urbano}. La gráfica 5 muestra los niveles de IPM, por sector en {municipio_seleccionado}:')
    doc.add_picture('grafico5.png', width=Inches(6.5))
    add_paragraph(doc, f'La principal problemática que impacta a una mayor proporción de la población en términos de pobreza multidimensional del municipio es el trabajo informal, a nivel municipal es de {trabajo_informal_municipio}, en el sector rural es de {trabajo_informal_rural} y en el sector urbano es de {trabajo_informal_urbano}.')
    doc.save(f'{municipio_seleccionado}.docx')

st.set_page_config(
    page_title="Departamento del Quindío",
    page_icon="DatosU/icon.png",  
    layout='wide',
    initial_sidebar_state="expanded"
)
def get_image_as_base64(file_path):
    with open(file_path, "rb") as file:
        data = file.read()
    return base64.b64encode(data).decode()

# Ruta de la imagen y conversión a base64
image_path = "DatosU/utp.png"
image_base64 = get_image_as_base64(image_path)

# Mostrar el título y la imagen en la misma línea usando HTML
st.markdown(
    f"""
    <div style="display: flex; align-items: center;">
        <h1 style="margin: 0; padding-right: 10px;">Departamento del Quindío</h1>
        <img src="data:image/png;base64,{image_base64}" alt="Logo" style="width: 50px;">
    </div>
    """,
    unsafe_allow_html=True
)

st.markdown(
    """
    <h2 style="font-size: 24px; color: #4CAF50;">
        Sembramos datos, cosechamos paz: la inteligencia artificial como el abono que fertiliza la justicia social
    </h2>
    """,
    unsafe_allow_html=True
)
shapefile_path = "DataBase/datos_agrarios/d_agrarios_quindio.shp"
dfData = gpd.read_file(shapefile_path)
dfAgro = gpd.read_file(shapefile_path)
dfQuindio = "DataBase/mapa_quindio/mapa_quindio.shp"
dfmapq = gpd.read_file(dfQuindio)
dfmapQ = gpd.read_file(dfQuindio)
dfGeneral = "DataBase/datos_quindio_generales.xlsx"
dfGeneral = pd.read_excel(dfGeneral)

precio_tierra = gpd.read_file('DatosU/DatosU/precio_tierra_quindio/precio_tierra_quindio.shp')

calidad_tierra = gpd.read_file('DatosU/DatosU/calidad_tierra_quindio/calidad_tierra_quindio.shp')


# Asegurarse de que las geometrías sean de tipo MultiPolygon o Polygon
dfData = dfData[dfData['geometry'].apply(lambda x: isinstance(x, (MultiPolygon, Polygon)))]

tab1,tab2,tab3,tab4,tab5=st.tabs(['Chatbot','Inicio','Mapa geográfico','Mapa cultivos','Reforma Agraria'])

with tab2:
    dfmapq['LATITUD'] = dfmapq.geometry.centroid.y  # Extraer la latitud de los centroides
    dfmapq['LONGITUD'] = dfmapq.geometry.centroid.x  # Extraer la longitud de los centroides

    # Crear un mapa coroplético usando Plotly
    fig = px.choropleth(dfmapq,
                        geojson=dfmapq.geometry,
                        locations=dfmapq.index,
                        color="AREA",  # Campo para colorear según el área
                        hover_name="MPIO_CNMBR",  # Nombre del municipio en el hover
                        hover_data={"LATITUD": True, "LONGITUD": True, "AREA": True, "MPIO_CDPMP": True},
                        title="Mapa del Quindío por Municipio",
                        labels={"AREA": "Área del Municipio"},
                        color_continuous_scale="Magma",
                        height=600)

    # Configurar el mapa para que solo muestre la región del Quindío
    fig.update_geos(fitbounds="locations", visible=False, projection_type="mercator")

    # Mostrar el mapa en Streamlit
    st.plotly_chart(fig, use_container_width=True)
    municipios = [
    'Armenia', 'Buenavista', 'Calarcá', 'Circasia', 'Córdoba', 'Filandia',
    'Génova', 'La Tebaida', 'Montenegro', 'Pijao', 'Quimbaya', 'Salento'
    ]

    # Crear columnas con proporciones ajustadas para mejor alineación
    col1, col2 = st.columns([4, 1])  # Ajusta los anchos de las columnas según sea necesario

    with col1:
        selected_municipio = st.selectbox("Selecciona un municipio del Quindío", municipios)

    with col2:
        # Agregar un espacio para alinear el botón
        st.write("")  # Espacio en blanco para alineación vertical
        if st.button("Generar informe"):            
            st.write(f"Has seleccionado el municipio de {selected_municipio}")
            generar_informes(selected_municipio)
# Crear el gráfico de Plotly con los centroides
with tab3: 
    dfmapq['LATITUD'] = dfmapq.geometry.centroid.y  # Extraer la latitud de los centroides
    dfmapq['LONGITUD'] = dfmapq.geometry.centroid.x
    
    dfmapq['MPIO_CDPMP'] = dfmapq['MPIO_CDPMP'].astype(int)  # Convertir a entero
    dfGeneral['Código Municipio'] = dfGeneral['Código Municipio'].astype(int)  # Convertir a entero  
    dfmapq = dfmapq.merge(dfGeneral, left_on='MPIO_CDPMP', right_on='Código Municipio', how='left')      
    
    # Crear columnas para los selectbox
    col1, col2, col3 = st.columns(3)
    
    opciones_estrato = ['Estrato 1', 'Estrato 2', 'Estrato 3', 'Estrato 4', 'Estrato 5', 'Estrato 6']
    opciones_sexo = ['Masculino', 'Femenino']
    opciones_escolaridad = ['Primaria', 'Secundaria', 'Universitaria', 'Posgrado']

    # Crear selectbox en cada columna
    with col1:
        campo_seleccionado_1 = st.selectbox("Selecciona el estrato", ["Todo"] + opciones_estrato)
    with col2:
        campo_seleccionado_2 = st.selectbox("Selecciona el sexo", ["Todo"] + opciones_sexo)
    with col3:
        campo_seleccionado_3 = st.selectbox("Selecciona el nivel de escolaridad", ["Todo"] + opciones_escolaridad)

    hover_data = {}

    # Copia del DataFrame para aplicar los filtros
    df_filtrado = dfmapq.copy()

    # Aplicar filtros según los valores seleccionados
    filter_conditions = []

    # Filtrar por estrato
    if campo_seleccionado_1 != "Todo":
        if campo_seleccionado_1 == "Estrato 1":
            filter_conditions.append(df_filtrado['Estrato 1'] > 0)
        elif campo_seleccionado_1 == "Estrato 2":
            filter_conditions.append(df_filtrado['Estrato 2'] > 0)
        elif campo_seleccionado_1 == "Estrato 3":
            filter_conditions.append(df_filtrado['Estrato 3'] > 0)
        elif campo_seleccionado_1 == "Estrato 4":
            filter_conditions.append(df_filtrado['Estrato 4'] > 0)
        elif campo_seleccionado_1 == "Estrato 5":
            filter_conditions.append(df_filtrado['Estrato 5'] > 0)
        elif campo_seleccionado_1 == "Estrato 6":
            filter_conditions.append(df_filtrado['Estrato 6'] > 0)

    # Filtrar por sexo
    if campo_seleccionado_2 != "Todo":
        if campo_seleccionado_2 == "Masculino":
            filter_conditions.append(df_filtrado['Hombres'] > 0)
        elif campo_seleccionado_2 == "Femenino":
            filter_conditions.append(df_filtrado['Mujeres'] > 0)

    # Filtrar por escolaridad
    if campo_seleccionado_3 != "Todo":
        if campo_seleccionado_3 == "Primaria":
            filter_conditions.append(df_filtrado['Primaria'] > 0)
        elif campo_seleccionado_3 == "Secundaria":
            filter_conditions.append(df_filtrado['Secundaria'] > 0)
        elif campo_seleccionado_3 == "Universitaria":
            filter_conditions.append(df_filtrado['Universitaria'] > 0)
        elif campo_seleccionado_3 == "Posgrado":
            filter_conditions.append(df_filtrado['Posgrado'] > 0)

    if filter_conditions:
        df_filtrado = df_filtrado[np.logical_and.reduce(filter_conditions)]
    else:
        # Si no se aplicó ningún filtro, mostrar el DataFrame completo
        df_filtrado = dfmapq
    divEst = (df_filtrado['Estrato 1'] + df_filtrado['Estrato 2'] + df_filtrado['Estrato 3'] + df_filtrado['Estrato 4'] + df_filtrado['Estrato 5'] + df_filtrado['Estrato 6'])   
    divSex = (df_filtrado['Hombres'] + df_filtrado['Mujeres'])
    divEsc = (df_filtrado['Primaria'] + df_filtrado['Secundaria'] + df_filtrado['Universitaria'] + df_filtrado['Posgrado'])
    
    # Calcular porcentajes
    df_filtrado['Estrato 1'] = ((df_filtrado['Estrato 1'] / divEst) * 100).round(1).astype(str) + '%'
    df_filtrado['Estrato 2'] = ((df_filtrado['Estrato 2'] / divEst) * 100).round(1).astype(str) + '%'
    df_filtrado['Estrato 3'] = ((df_filtrado['Estrato 3'] / divEst) * 100).round(1).astype(str) + '%'
    df_filtrado['Estrato 4'] = ((df_filtrado['Estrato 4'] / divEst) * 100).round(1).astype(str) + '%'
    df_filtrado['Estrato 5'] = ((df_filtrado['Estrato 5'] / divEst) * 100).round(1).astype(str) + '%'
    df_filtrado['Estrato 6'] = ((df_filtrado['Estrato 6'] / divEst) * 100).round(1).astype(str) + '%'
    df_filtrado['Hombres'] = ((df_filtrado['Hombres'] / divSex) * 100).round(1).astype(str) + '%'
    df_filtrado['Mujeres'] = ((df_filtrado['Mujeres'] / divSex) * 100).round(1).astype(str) + '%'
    df_filtrado['Primaria'] = ((df_filtrado['Primaria'] / divEsc) * 100).round(1).astype(str) + '%'
    df_filtrado['Secundaria'] = ((df_filtrado['Secundaria'] / divEsc) * 100).round(1).astype(str) + '%'
    df_filtrado['Universitaria'] = ((df_filtrado['Universitaria'] / divEsc) * 100).round(1).astype(str) + '%'
    df_filtrado['Posgrado'] = ((df_filtrado['Posgrado'] / divEsc) * 100).round(1).astype(str) + '%'
        
        
    if campo_seleccionado_1 == "Estrato 1":
        hover_data['Estrato 1'] = True
    elif campo_seleccionado_1 == "Estrato 2":
        hover_data['Estrato 2'] = True
    elif campo_seleccionado_1 == "Estrato 3":
        hover_data['Estrato 3'] = True
    elif campo_seleccionado_1 == "Estrato 4":
        hover_data['Estrato 4'] = True
    elif campo_seleccionado_1 == "Estrato 5":
        hover_data['Estrato 5'] = True
    elif campo_seleccionado_1 == "Estrato 6":
        hover_data['Estrato 6'] = True
    else:
        hover_data['Estrato 1'] = True
        hover_data['Estrato 2'] = True
        hover_data['Estrato 3'] = True
        hover_data['Estrato 4'] = True
        hover_data['Estrato 5'] = True
        hover_data['Estrato 6'] = True

    if campo_seleccionado_2 == "Masculino":
        hover_data['Hombres'] = True
    elif campo_seleccionado_2 == "Femenino":
        hover_data['Mujeres'] = True
    else:
        hover_data['Hombres'] = True
        hover_data['Mujeres'] = True
        
    if campo_seleccionado_3 == "Primaria":
        hover_data['Primaria'] = True
    elif campo_seleccionado_3 == "Secundaria":
        hover_data['Secundaria'] = True
    elif campo_seleccionado_3 == "Universitaria":
        hover_data['Universitaria'] = True
    elif campo_seleccionado_3 == "Posgrado":
        hover_data['Posgrado'] = True    
    else:
        hover_data['Primaria'] = True
        hover_data['Secundaria'] = True
        hover_data['Universitaria'] = True
        hover_data['Posgrado'] = True
    hover_data['AREA'] = True
    

    # Crear el gráfico con hover_data ajustado
    fig = px.choropleth(
        df_filtrado,
        geojson=df_filtrado.geometry,
        locations=df_filtrado.index,
        color='MPIO_CNMBR',
        hover_name="MPIO_CNMBR",
        hover_data= hover_data,
        labels={"AREA": "Área del Municipio"},
        color_continuous_scale="Magma",
        height=600
    )

    fig.update_geos(fitbounds="locations", visible=False, projection_type="mercator")
    
    # Mostrar el gráfico en Streamlit
    st.plotly_chart(fig, use_container_width=True, showlegend=True)
    
    # Crear gráficos de barras horizontales
    # Reorganizar el DataFrame para Estrato, Sexo y Escolaridad
    df_estrato = df_filtrado.melt(id_vars='MPIO_CNMBR', 
                                value_vars=['Estrato 1', 'Estrato 2', 'Estrato 3', 'Estrato 4', 'Estrato 5', 'Estrato 6'],
                                var_name='Estrato', value_name='Porcentaje')

    df_sexo = df_filtrado.melt(id_vars='MPIO_CNMBR', 
                            value_vars=['Hombres', 'Mujeres'], 
                            var_name='Sexo', value_name='Porcentaje')

    df_escolaridad = df_filtrado.melt(id_vars='MPIO_CNMBR', 
                                    value_vars=['Primaria', 'Secundaria', 'Universitaria', 'Posgrado'], 
                                    var_name='Escolaridad', value_name='Porcentaje')

    # Graficar Estrato
    fig_estrato = px.bar(
        df_estrato,  # DataFrame reorganizado para Estrato
        y='MPIO_CNMBR',  # Municipios en el eje Y
        x='Porcentaje',  # Porcentaje en el eje X
        color='Estrato',  # Colorear las barras por el Estrato
        orientation='h',  # Barras horizontales
        title='Distribución por Estrato',
        labels={'MPIO_CNMBR': 'Municipios', 'Porcentaje': 'Porcentaje (%)'}
    )

    # Graficar Sexo
    fig_sexo = px.bar(
        df_sexo,  # DataFrame reorganizado para Sexo
        y='MPIO_CNMBR',  # Municipios en el eje Y
        x='Porcentaje',  # Porcentaje en el eje X
        color='Sexo',  # Colorear las barras por Sexo
        orientation='h',  # Barras horizontales
        title='Distribución por Sexo',
        labels={'MPIO_CNMBR': 'Municipios', 'Porcentaje': 'Porcentaje (%)'}
    )

    # Graficar Escolaridad
    fig_escolaridad = px.bar(
        df_escolaridad,  # DataFrame reorganizado para Escolaridad
        y='MPIO_CNMBR',  # Municipios en el eje Y
        x='Porcentaje',  # Porcentaje en el eje X
        color='Escolaridad',  # Colorear las barras por Escolaridad
        orientation='h',  # Barras horizontales
        title='Distribución por Escolaridad',
        labels={'MPIO_CNMBR': 'Municipios', 'Porcentaje': 'Porcentaje (%)'}
    )

    # Mostrar gráficos en Streamlit
    st.plotly_chart(fig_estrato)
    st.plotly_chart(fig_sexo)
    st.plotly_chart(fig_escolaridad)
    # st.write(df_filtrado.head())  # Muestra las primeras filas del DataFrame para verificar su estructura

        

with tab4:
    # Asegúrate de que 'geometry' sea la columna de coordenadas
    dfAgro = dfAgro.to_crs(epsg=4326)  # Convertir a latitud y longitud si es necesario
    dfmapQ = dfmapQ.to_crs(epsg=4326)
    dfmapQ['LATITUD'] = dfmapQ.geometry.centroid.y  # Extraer la latitud de los centroides
    dfmapQ['LONGITUD'] = dfmapQ.geometry.centroid.x
    # Esto crea un ciclo repetido de cultivos hasta que tenga la misma longitud que dfAgro
    cultivos = ['aguacate', 'banano', 'café', 'plátano']
    n_repeticiones = len(dfAgro) // len(cultivos) + 1  # Asegurarnos de que sea suficiente
    variables = (cultivos * n_repeticiones)[:len(dfAgro)]  # Recortar al tamaño exacto

    df_puntos = pd.DataFrame({
        'variable': variables,
        'lat': dfAgro.geometry.centroid.y,  # Latitudes de los centroides
        'lon': dfAgro.geometry.centroid.x,  # Longitudes de los centroides
    })

    # Crear el gráfico choropleth
    fig = px.choropleth(
                            df_filtrado,
                            geojson=df_filtrado.geometry.__geo_interface__,  # Asegúrate de que 'geometry' esté bien configurado
                            locations=df_filtrado.index,
                            color='MPIO_CNMBR',  # Aquí puedes usar cualquier columna que represente valores numéricos o categóricos
                            hover_name="MPIO_CNMBR",  # Aquí debería ir el nombre del municipio o una columna relevante
                            hover_data=hover_data,  # Datos adicionales que aparecerán al pasar el cursor
                            labels={"AREA": "Área del Municipio"},  # Etiquetas para los datos
                            color_continuous_scale="Magma",  # Escala de colores continua
                            height=600
                        )


    # Configurar el mapa para que muestre los cultivos
    fig.update_geos(fitbounds="locations", visible=False, projection_type="mercator")

    # Crear el scatter_geo para los puntos (sin mostrar puntos por defecto)
    scatter = go.Scattergeo(
        lon=[],
        lat=[],
        text=[],
        mode='markers',
        marker=dict(size=8, color='red', opacity=0.7),
        showlegend=False
    )

    # Radio buttons para seleccionar la variable de cultivo
    radio_var = st.radio("Selecciona un cultivo:", ('Aguacate', 'Banano', 'Café', 'Plátano'),horizontal=True)

    # Filtrar los puntos según la selección de la variable
    color_map = {'Aguacate': 'darkgreen', 'Banano': 'yellow', 'Café': 'brown', 'Plátano': 'green'}
    if radio_var == 'Aguacate':
        selected_points = dfAgro[dfAgro['ap_agaucat'] == 1]
    elif radio_var == 'Banano':
        selected_points = dfAgro[dfAgro['ac_banano'] == 1]
    elif radio_var == 'Café':
        selected_points = dfAgro[dfAgro['ac_cafe'] == 1]
    else:
        selected_points = dfAgro[dfAgro['ac_platano'] == 1]
        
    # Obtener el color correspondiente
    selected_color = color_map.get(radio_var, 'gray')

    # Agregar los puntos seleccionados al gráfico
    scatter = go.Scattergeo(
        lon=selected_points.geometry.centroid.x,  # Longitudes de los centroides
        lat=selected_points.geometry.centroid.y,  # Latitudes de los centroides
        text=selected_points['municipio'],  # Información adicional sobre el punto, como el nombre del municipio
        mode='markers',
        marker=dict(size=8, color= selected_color, opacity=0.7),  # Estilo de los puntos
        showlegend=False
    )

    # Añadir scatter_plot al gráfico
    fig.add_trace(scatter)

    # Configurar la vista del mapa
    fig.update_geos(fitbounds="locations", visible=False)

    # Mostrar el gráfico en Streamlit
    st.plotly_chart(fig, use_container_width=True)
    
with tab5:
    # Asegurarse de que está en EPSG:4326
    precio_tierra = precio_tierra.to_crs("EPSG:4326")
    calidad_tierra = calidad_tierra.to_crs("EPSG:4326")

    # Reproyectar a un CRS proyectado para calcular los centroides
    precio_tierra_proj = precio_tierra.to_crs("EPSG:3857")
    calidad_tierra_proj = calidad_tierra.to_crs("EPSG:3857")

    # Calcular los centroides
    precio_tierra_centroid = precio_tierra_proj.geometry.centroid.to_crs("EPSG:4326")
    calidad_tierra_centroid = calidad_tierra_proj.geometry.centroid.to_crs("EPSG:4326")

    # Convertir a GeoJSON
    geojson_data_precio = json.loads(precio_tierra.to_json())
    geojson_data_calidad = json.loads(calidad_tierra.to_json())

    # 1. Crear el diccionario con el orden personalizado
    orden_personalizado = {
        'Mayor a 1 - hasta 2': 1,
        'Mayor a 3 - hasta 5': 2,
        'Mayor a 8 - hasta 10': 3,
        'Mayor a 10 - hasta 15': 4,
        'Mayor a 15 - hasta 20': 5,
        'Mayor a 20 - hasta 30': 6,
        'Mayor a 30 - hasta 40': 7,
        'Mayor a 40 - hasta 50': 8,
        'Mayor a 50 - hasta 60': 9,
        'Mayor a 60 - hasta 70': 10,
        'Mayor a 70 - hasta 80': 11,
        'Mayor a 80 - hasta 100': 12,
        'Mayor a 100 - hasta 120': 13,
        'Mayor a 120 - hasta 140': 14,
        'Mayor a 140 - hasta 160': 15,
        'Mayor a 180 - hasta 200': 16,
        'Mayor a 550 - hasta 600': 17,
        'Mayor a 650 - hasta 700': 18
    }

    # Asignar el orden personalizado a la columna 'rango_prec' en el DataFrame del mapa
    precio_tierra['orden'] = precio_tierra['rango_prec'].map(orden_personalizado)

    # Función para crear el gráfico de barras
    def crear_grafico_barras(df):
        fig = px.bar(
            data_frame=df,
            x='hectareas',
            y='rango_prec',
            orientation='h',
            labels={'hectareas': 'Extensión total (ha)', 'rango_prec': 'Rango de precio (millones de pesos por ha)'},
            color='orden',  # Usar la columna 'orden' para colorear las barras
            color_continuous_scale="YlGnBu"  # Aseguramos que la escala de colores sea Viridis
        )
        fig.update_layout(margin={"r": 0, "t": 0, "l": 0, "b": 0})
        fig.update_yaxes(
            categoryorder='array',
            categoryarray=df['rango_prec']
        )
        fig.update_layout(
            margin={"l": 200, "r": 10, "t": 10, "b": 10},  # Aumentamos el margen izquierdo
        )
        fig.update_yaxes(
            automargin=True,
            tickfont=dict(size=10),  # Reducimos el tamaño de la fuente si es necesario
        )
        return fig

    # Función para crear el mapa de precios
    def crear_mapa_precios(df, geojson_data):
        fig = px.choropleth_mapbox(
            df,
            geojson=geojson_data,
            locations=df.index,
            color="orden",  # Usar la columna 'orden' para colorear el mapa
            color_continuous_scale="YlGnBu",
            mapbox_style="carto-positron",
            center={"lat": precio_tierra_centroid.y.mean(), "lon": precio_tierra_centroid.x.mean()},
            zoom=8.5,
            opacity=0.6,
            hover_name="rango_prec",
            hover_data={"orden": False, "rango_prec": True}
        )
        fig.update_traces(marker_line_width=0.2)
        fig.update_traces(hovertemplate="<b>%{hovertext}</b><br>Millones de pesos por hectárea")
        fig.update_layout(margin={"r": 0, "t": 0, "l": 0, "b": 0})
        return fig

    # Agrupar y sumar las hectáreas por rango de precio
    df_agrupado = precio_tierra.groupby('rango_prec')['hectareas'].sum().reset_index()

    # Asignar el orden a cada categoría en el DataFrame
    df_agrupado['orden'] = df_agrupado['rango_prec'].map(orden_personalizado)

    # Ordenar el DataFrame según el orden personalizado
    df_agrupado = df_agrupado.sort_values('orden', ascending=False)

    # Crear los gráficos
    fig_barra = crear_grafico_barras(df_agrupado)
    fig_precio = crear_mapa_precios(precio_tierra, geojson_data_precio)

    # Combinar ambos gráficos en un único layout
    fig = make_subplots(
        rows=1, cols=2,
        column_widths=[0.5, 0.5],  
        specs=[[{"type": "xy"}, {"type": "mapbox"}]]  
    )

    # Añadir el gráfico de barras al subplot izquierdo
    for trace in fig_barra.data:
        fig.add_trace(trace, row=1, col=1)

    # Añadir el mapa al subplot derecho
    for trace in fig_precio.data:
        fig.add_trace(trace, row=1, col=2)

    # Actualizar layout general
    fig.update_layout(
        mapbox=dict(
            style="carto-positron",
            domain={'x': [0.5, 1], 'y': [0, 1]},  
            center={"lat": precio_tierra_centroid.y.mean(), "lon": precio_tierra_centroid.x.mean()},
            zoom=8.5
        ),
        showlegend=False,
    )

    # Ajustar los márgenes generales de la figura
    fig.update_layout(margin={"r": 0, "t": 0, "l": 0, "b": 0})

    # Mostrar gráfico combinado
    # fig.show()
    
    # Configurar el mapa de calidad UFH
    def crear_mapa_calidad(df, geojson_data):
        fig = px.choropleth_mapbox(
            df,
            geojson=geojson_data,
            locations=df.index,
            color="Valor_UFH",
            color_continuous_scale="YlGn",
            mapbox_style="carto-positron",
            center={"lat": calidad_tierra_centroid.y.mean(), "lon": calidad_tierra_centroid.x.mean()},
            zoom=8.5,
            opacity=0.6,
            hover_name="Valor_UFH",
            hover_data={"Valor_UFH": False}
        )
        fig.update_traces(marker_line_width=0.2)
        fig.update_traces(hovertemplate="<b>%{hovertext}</b><br>Nivel de calidad de la Unidad Física Homogénea (UFH) 1 es la mejor 13 la más restrictiva")
        fig.update_layout(margin={"r":0,"t":0,"l":0,"b":0})
        return fig

    # Crear gráfico de barras horizontales usando los datos de calidad UFH
    def crear_grafico_barras_calidad(df):
        fig = px.bar(
            data_frame=df,
            x='hectareas',
            y='Valor_UFH',
            orientation='h',
            labels={'hectareas': 'Extensión total (ha)', 'Valor_UFH': 'Nivel de calidad UFH'},
            color='Valor_UFH',  
            color_continuous_scale="YlGn" 
        )
        fig.update_layout(margin={"r": 0, "t": 0, "l": 0, "b": 0})
        fig.update_yaxes(
            categoryorder='array',
            categoryarray=df['Valor_UFH'],
        )
        fig.update_layout(
            margin={"l": 200, "r": 10, "t": 10, "b": 10}, 
        )
        fig.update_yaxes(
            automargin=True,
            tickfont=dict(size=10),  
        )
        return fig

    # Agrupar y sumar las hectáreas por nivel de calidad UFH
    df_calidad_agrupado = calidad_tierra.groupby('Valor_UFH')['hectareas'].sum().reset_index()

    # Ordenar el DataFrame por 'Valor_UFH'
    df_calidad_agrupado = df_calidad_agrupado.sort_values('Valor_UFH')

    # Crear los gráficos
    fig_barra_calidad = crear_grafico_barras_calidad(df_calidad_agrupado)
    fig_calidad = crear_mapa_calidad(calidad_tierra, geojson_data_calidad)

    # Combinar ambos gráficos en un único layout
    # Streamlit Layout
    # st.title("Visualización de Precio y Calidad de Tierra")
    
    tab6, tab7, tab8 = st.tabs(["Precio y calidad de Tierra", "Análisis Territorial", "Validación de resultados"])

    with tab6:
        col1, col2 = st.columns(2)  # Crear dos columnas        
        #     # Título en la columna 2
        #     st.subheader("")            
           
        with col1:
            # Primera fila de gráficos con texto al lado
            st.plotly_chart(fig_barra, use_container_width=True)
            st.write("Cantidad de hectáreas por rango de precio.")

            st.plotly_chart(fig_precio, use_container_width=True)
            st.write("Districución de rangos de precio de la tierra en el departamento del Quindío.")
            
            st.write("""
            <div style="text-align: justify;">
                Utilizando como fuente el dataset precios comerciales de la tierra rural agropecuaria del portal Datos Abiertos se filtró la información correspondiente al departamento del Quindío en formato Shapefile, el cual contiene la distribución por rangos de precio de la tierra como se observa en la gráfica anterior. Se puede observar que las tierras al sur y al oriente del departamento son las que tienen los menores rangos de precios, por el contrario, las tierras que rodean el centro urbano de Armenia son las más costosas. La zona noroccidental de Quindío presenta tierras con valores medios, esto debido a su cercanía a los centros urbanos, municipios y ciudades de otros departamentos como Pereira y Cartago.
            </div>
            """, unsafe_allow_html=True)

        with col2:
            # Segunda fila de gráficos con texto al lado
            st.plotly_chart(fig_calidad, use_container_width=True)
            st.write("Unidades del suelo del Quindío en función de su restricción productiva.")

            st.plotly_chart(fig_barra_calidad, use_container_width=True)
            st.write("Cantidad de hectáreas por unidad de suelo.")
            
            st.write("""
            <div style="text-align: justify;">
                Utilizando como fuente el dataset unidades físicas homogéneas en Colombia para el cálculo de la Unidad Agrícola Familiar -UAF- del portal Datos Abiertos se filtró la información correspondiente al departamento del Quindío en formato Shapefile, el cual contiene la distribución por Unidades Físicas Homogéneas (restricción para la productividad en la tierra, siendo 13 la más restrictiva y 1 la menos restrictiva) como se observa en la gráfica anterior. Se evidencia que la zona sur y oriente del Quindío presentan altas restricciones para la productividad, principalmente asociadas a áreas de protección ambiental. Las zonas que presentan una menor restricción en su productividad se encuentran en la zona noroccidental entre municipios de Armenia, Montenegro, Circasia, Filandia y Quimbaya.                
            </div>
            """, unsafe_allow_html=True)
            
            

    with tab7:
        # Cargar los datos desde el archivo Excel
        socioeconomico_quindio = pd.read_excel('DatosU/DatosU/Socioeconomico_Quindio.xlsx')

        # Filtrar los datos para el año 2018 y seleccionar las columnas deseadas
        datos_2018 = socioeconomico_quindio[socioeconomico_quindio['Año'] == 2018][['Municipio', 'Población Urbana', 'Población Rural']]

        # Calcular los porcentajes
        datos_2018['Total Población'] = datos_2018['Población Urbana'] + datos_2018['Población Rural']
        datos_2018['% Población Urbana'] = (datos_2018['Población Urbana'] / datos_2018['Total Población']) * 100
        datos_2018['% Población Rural'] = (datos_2018['Población Rural'] / datos_2018['Total Población']) * 100

        # Formatear los porcentajes para hover
        datos_2018['% Población Urbana (hover)'] = datos_2018['% Población Urbana'].map("{:.1f}%".format)
        datos_2018['% Población Rural (hover)'] = datos_2018['% Población Rural'].map("{:.1f}%".format)

        # Configurar la página en Streamlit
        st.title("Distribución de población en el Quindío (2018)")
        st.write("""
            <div style="text-align: justify;">
                Del análisis en de los mapas relacionados anterior se observa que las tierras que guardan una mejor relación representada en unas menores restricciones productivas de sus Unidades Físicas Homogéneas y precios de la tierra moderados y accesibles para su compra en cantidades suficientes por parte del Gobierno Nacional son aquellas que se encuentran que al noroccidente en los municipios de Filandia, Circasia, Quimbaya, Montenegro y La Tebaida. Estas deben ser el foco de función de la producción agropecuaria para garantizar seguridad alimentaria y alternativas económicas a las poblaciones rurales del Quindío.

\nSin embargo, las tierras al occidente del Quindío muestran una alta correlación entre sus elevados niveles de restricción, debido a la presencia de áreas protegidas para la conservación ambiental, y su bajo precio por hectárea. A pesar de esto, como se puede observar en las gráficas siguientes, estas no son tierras despobladas; por el contrario, están habitadas por poblaciones establecidas, muchas de ellas en el sector rural, que también requieren el apoyo del Estado.
            </div>
            """, unsafe_allow_html=True)
        

        # Gráfico interactivo para porcentaje de Población Urbana
        fig_urbana = px.bar(
            datos_2018,
            x='% Población Urbana',
            y='Municipio',
            orientation='h',  # Gráfico horizontal
            title='Porcentaje de población urbana en el Quindío (2018)',
            labels={'% Población Urbana': 'Porcentaje de Población Urbana', 'Municipio': 'Municipios'},
            color='% Población Urbana',  # Color por cantidad
            hover_data={'% Población Urbana (hover)': True, 'Municipio': False}  # Mostrar solo el porcentaje con formato en hover
        )
        fig_urbana.update_layout(xaxis_title='Porcentaje (%)', yaxis_title='Municipios')
        st.plotly_chart(fig_urbana, use_container_width=True)

        # Gráfico interactivo para porcentaje de Población Rural
        fig_rural = px.bar(
            datos_2018,
            x='% Población Rural',
            y='Municipio',
            orientation='h',  # Gráfico horizontal
            title='Porcentaje de población rural en el Quindío (2018)',
            labels={'% Población Rural': 'Porcentaje de Población Rural', 'Municipio': 'Municipios'},
            color='% Población Rural',  # Color por cantidad
            hover_data={'% Población Rural (hover)': True, 'Municipio': False}  # Mostrar solo el porcentaje con formato en hover
        )
        fig_rural.update_layout(xaxis_title='Porcentaje (%)', yaxis_title='Municipios')
        st.plotly_chart(fig_rural, use_container_width=True)

        # Gráfico combinado: Porcentajes de Población Urbana y Rural
        fig_combinado = go.Figure()

        # Agregar Porcentaje de Población Urbana
        fig_combinado.add_trace(go.Bar(
            y=datos_2018['Municipio'],
            x=datos_2018['% Población Urbana'],
            name='% Población Urbana',
            orientation='h',
            marker=dict(color='skyblue'),
            hovertemplate='%{x:.1f}%'
        ))

        # Agregar Porcentaje de Población Rural
        fig_combinado.add_trace(go.Bar(
            y=datos_2018['Municipio'],
            x=datos_2018['% Población Rural'],
            name='% Población Rural',
            orientation='h',
            marker=dict(color='cornflowerblue'),
            hovertemplate='%{x:.1f}%'
        ))

        # Configurar diseño
        fig_combinado.update_layout(
            title='Porcentaje de población urbana y rural en el Quindío (2018)',
            barmode='stack',  # Barras apiladas
            xaxis_title='Porcentaje (%)',
            yaxis_title='Municipios',
            legend_title='Tipo de Población',
            height=600
        )

        # Mostrar gráfico combinado en Streamlit
        st.plotly_chart(fig_combinado, use_container_width=True)
        st.write("""
            <div style="text-align: justify;">
                En este contexto, el aplicativo se presenta como una herramienta clave para la toma de decisiones de los funcionarios públicos. Gracias a la identificación visual que ofrecen los mapas y al cruce de información entre las zonas de interés ambiental y las áreas con precios de tierras bajos, se facilita la adquisición de predios para la creación de corredores biológicos, zonas de recarga de acuíferos y áreas aferentes a los acueductos comunitarios y municipales, entre otras.

Sin embargo, no es posible comprar todos los predios, ni es viable reubicar a toda la población campesina presente en la zona. Por lo que el aplicativo también se presta como herramienta para la orientar decisiones hacia la promoción de prácticas de producción y consumo sostenible respaldas por entidades como las Alcaldías Municipales y la Corporación Autónoma que garantice la calidad de vida las comunidades rurales en la zona.
            </div>
            """, unsafe_allow_html=True)
        
        with tab8:
            st.write("""
            <div style="text-align: justify;">
                Los resultados obtenidos a partir de los análisis fueron validados mediante modelos de Machine Learning, específicamente con algoritmos de clasificación como Random Forest y Gradient Boosting, los cuales alcanzaron métricas de accuracy del 75% en la estimación de los precios de la tierra. Esto demuestra que las relaciones entre datos socioeconómicos, agrarios y territoriales constituyen una estrategia efectiva para la toma de decisiones, especialmente en la identificación de tierras con mayor viabilidad para su adquisición a través de la Agencia Nacional de Tierras y otras entidades involucradas en la Reforma Agraria.

Concretamente, los modelos identificaron como las principales variables a considerar en la adquisición de tierras: la ubicación de los municipios, los tipos de suelos, el clima, la profundidad radicular, el perfil de los suelos y las restricciones propias de las Unidades Físicas Homogéneas. Además, este modelo contribuirá a agilizar los procesos de análisis, proporcionando a los funcionarios públicos, técnicos y comunidades un insumo respaldado tanto por estadísticas como por metodologías sólidas, que estará listo para ser reinterpretado y ajustado según las realidades territoriales construidas a partir de los espacios de participación.

El modelo puede ser consultado a través del Jupyter Notebook denominado “modelo_predicción_reforma_agraria” en el repositorio de GitHub del proyecto.
            """, unsafe_allow_html=True)

        
        # st.subheader("Calidad de Tierra")
        # Mostrar gráficos de forma horizontal
        # col1, col2 = st.columns(2)  # Crear dos columnas
        # with col1:
        #     st.write("Aquí se visualizan las barras que combinan calidad y precio para una mejor comparación.")
            
        # with col2:
        #     st.write("Aquí se visualizan las barras que combinan calidad y precio para una mejor comparación.")
            
    
with tab1:
    # Descargar los recursos necesarios de NLTK si no los tienes
    nltk.download('punkt')

    # Inicializar el cliente de Groq
    with open('config.json') as config_file:
        config = json.load(config_file)

    # Inicializar el cliente con la API key del JSON
    client = Groq(api_key=config["GROQ_API_KEY"])

    # Función para cargar datos de archivos .xlsx
    def load_xlsx(file_path):
        return pd.read_excel(file_path)

    # Función para obtener la respuesta de la IA
    def get_ai_response(messages):
        completion = client.chat.completions.create(
            model="llama-3.1-70b-versatile",
            messages=messages,
            temperature=0.2,
            max_tokens=512,
            stream=True,
        )
        response = "".join(chunk.choices[0].delta.content or "" for chunk in completion)
        return response

    # Función para buscar en los datos cargados
    def search_in_data(query, dataframe):
        result = []
        
        # Asegurarse de que el query no sea vacío
        if not query.strip():
            return result
        
        # Buscar el query en todas las columnas de la tabla
        for column in dataframe.columns:
            # Si la columna es de tipo string (objeto), usar str.contains
            if dataframe[column].dtype == 'object':
                # Filtramos las filas donde la columna no es NaN y aplicamos str.contains
                valid_rows = dataframe[column].dropna()
                result += dataframe[valid_rows.str.contains(query, case=False, na=False)].values.tolist()
            
            # Si la columna es de tipo numérico (int o float), buscar coincidencias numéricas
            elif pd.api.types.is_numeric_dtype(dataframe[column]):
                # Verificar si la consulta es un número y buscarlo en la columna
                try:
                    numeric_query = float(query)  # Intentamos convertir el query a número
                    matching_rows = dataframe[dataframe[column] == numeric_query]
                    result += matching_rows.values.tolist()  # Añadimos los resultados encontrados
                except ValueError:
                    # Si no se puede convertir el query a número, simplemente lo ignoramos
                    pass
        
        return result

    # Función para procesar el query (minúsculas y tokenización)
    def process_query(query):
        query_lower = query.lower()  # Convertimos a minúsculas
        try:
            query_tokens = word_tokenize(query_lower)  # Tokenizamos el query
        except Exception as e:
            print(f"Error al tokenizar: {e}")
            query_tokens = query_lower.split()  # Si falla, dividimos en espacios
        return " ".join(query_tokens)  # Retornamos el query tokenizado

    # Función para reiniciar el chat al cambiar de categoría
    def reiniciar_chat():
        """Reinicia el historial del chat cuando se selecciona una nueva categoría."""
        st.session_state['messages'] = []  # Limpiar el historial de mensajes
        st.session_state['dataframe'] = None  # Limpiar el dataframe
        st.session_state['current_category'] = None  # Limpiar la categoría actual

    # Función principal del chat
    def chat():
        # st.title("Información departamento del Quindío")
        st.write("Bienvenido a tu asistente de información de tu departamento.")

        if 'messages' not in st.session_state:
            st.session_state['messages'] = []

        # Categorías y archivos
        categories = {
            "Cultivos": "DatosU\DatosU\EVAS_Quindio.xlsx",
            "Población Vulnerable": "DatosU\DatosU\IPM_Quindio.xlsx",
            "Datos socioeconómicos": "DatosU\DatosU\Socioeconomico_Quindio.xlsx",
            "Información general": "DatosU\DatosU\datos_quindio_generales.xlsx"
        }

        # Selector de categoría
        category = st.selectbox("Selecciona la categoría:", options=list(categories.keys()))

        # Revisar si se seleccionó una nueva categoría
        if 'current_category' not in st.session_state or st.session_state['current_category'] != category:
            # Reiniciar el chat y cargar los datos correspondientes a la nueva categoría
            reiniciar_chat()

            # Cargar el archivo correspondiente basado en la categoría seleccionada
            dataframe = load_xlsx(categories[category])
            st.session_state['dataframe'] = dataframe  # Guardar el dataframe en la sesión
            st.session_state['current_category'] = category  # Guardar la categoría seleccionada

            # Imprimir en consola el cambio de dataset
            print(f"Se ha cargado el dataset para la categoría '{category}' desde el archivo: {categories[category]}")

            # Mensaje de sistema (no se muestra en el chat, solo se usa para el modelo)
            promptSistema = f"""
            Eres un asistente de datos del departamento del Quindío, especializado en: {category}.
            Tu tarea es responder de manera directa y concisa, sin agregar información innecesaria o redundante.
            Si la pregunta puede ser respondida con los datos cargados en el archivo, responde con esos datos. Si no, puedes usar tu conocimiento general.
            """
            st.session_state['messages'].append({"role": "system", "content": promptSistema})

        # Obtener el dataframe desde la sesión
        dataframe = st.session_state.get('dataframe', None)
        
        # Función para enviar el mensaje
        def submit():
            user_input = st.session_state.user_input
            if user_input.lower() == 'salir':
                st.write("¡Gracias por chatear!")
                st.stop()

            # Procesar el query (convertir a minúsculas y tokenizar)
            processed_query = process_query(user_input)

            # Buscar en los datos del archivo seleccionado primero
            if dataframe is not None:
                search_results = search_in_data(processed_query, dataframe)

                if search_results:
                    # Si se encuentra información relevante en los archivos, responder con los resultados
                    response = f"Encontré estos resultados en los datos del archivo:\n{search_results}"
                    st.session_state['messages'].append({"role": "assistant", "content": response})
                else:
                    # Si no se encuentra en los datos, pasar la consulta al modelo LLM
                    st.session_state['messages'].append({"role": "user", "content": user_input})
                    with st.spinner("Buscando información..."):
                        ai_response = get_ai_response(st.session_state['messages'])
                        # Solo agregar la respuesta del modelo si no se encontró nada en los datos
                        response = ai_response
                        st.session_state['messages'].append({"role": "assistant", "content": ai_response})
            else:
                response = "No se ha cargado ningún archivo para la categoría seleccionada."
                st.session_state['messages'].append({"role": "assistant", "content": response})

            # Limpiar la entrada del usuario
            st.session_state.user_input = ""

        # Mostrar mensajes anteriores
        for message in st.session_state['messages']:
            if message["role"] != "system":  # No mostrar el mensaje del sistema
                role = "Tú" if message["role"] == "user" else "Bot"
                icon = "🤖" if message["role"] == "assistant" else "💬"
                st.write(f"**{role} {icon}:** {message['content']}")  # Mostrar íconos con el rol
        
        # Formulario de entrada de usuario
        with st.form(key='chat_form', clear_on_submit=True):
            st.text_input("Tú:", key="user_input")
            submit_button = st.form_submit_button(label='Enviar', on_click=submit)

    # Ejecución principal
    if __name__ == "__main__":
        chat()
    




# streamlit run e:/Datos_U/visualAPP.py este es el comando para iniciar el aplicativo

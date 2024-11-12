#Importamos las librerias
import pandas as pd
import geopandas as gpd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

#importamos los archivos
mapa_base = gpd.read_file('DatosU\DatosU\mapa_quindio\mapa_quindio.shp')
evas_quindio = pd.read_excel('DatosU\DatosU\EVAS_Quindio.xlsx')
ipm_quindio = pd.read_excel('DatosU\DatosU\IPM_Quindio.xlsx')
general_quindio = pd.read_excel('DatosU\DatosU\datos_quindio_generales.xlsx')
socioeconomico_quindio = pd.read_excel('DatosU\DatosU\Socioeconomico_Quindio.xlsx')

print("mapa_base",mapa_base.columns)
print("evas_quindio",evas_quindio.columns)
print("ipm_quindio",ipm_quindio.columns)
print("general_quindio",general_quindio.columns)
print("socioeconomico_quindio",socioeconomico_quindio.columns)

#Seleccionamos el municipio
municipio_seleccionado = str(input('Ingrese el municipio que desea consultar: '))

#Ubicación de los municipios en el Departamento del Quindío
ubicaciones = {
    'Armenia': 'centro',
    'Buenavista': 'suroccidente',
    'Calarcá': 'oriente',
    'Circasia': 'centro-norte',
    'Córdoba': 'suroriente',
    'Filandia': 'norte',
    'Génova': 'sur',
    'La Tebaida': 'occidente',
    'Montenegro': 'occidente',
    'Pijao': 'sur',
    'Quimbaya': 'noroccidente',
    'Salento': 'nororiente',
    }

#### INTRODUCCIÓN ####


#Definir las variables de la introducción
#Extensión
datos_generales_quindio = general_quindio[general_quindio['Municipio'] == municipio_seleccionado]
extension = datos_generales_quindio['Área en Km2'].values[0]

#Población total 2022

dato_población_quindio = socioeconomico_quindio[(socioeconomico_quindio['Municipio'] == municipio_seleccionado)
                                                    & (socioeconomico_quindio['Año'] == 2022)]

poblacion_total = dato_población_quindio['Población Total'].values[0]

#Crear la figura y los ejes
plt.figure(figsize=(12,10))
ax = plt.axes()
mapa_base.plot(ax=ax, color='white', edgecolor='#363636')

#Filtrar el municipio de la variable de entrada
nombre_municipio = mapa_base[mapa_base['MPIO_CNMBR'] == municipio_seleccionado]

#Agregar el municipio en el mapa
nombre_municipio.plot(ax=ax, color='green', edgecolor='black')

#Agregar los detalles estéticos
plt.title(f'Municipio de {municipio_seleccionado}', fontsize=20, fontweight='bold')
ax.set_xlabel('Longitud', color='#363636', fontsize=14)
ax.set_ylabel('Latitud', color='#363636', fontsize=14)
plt.grid(color='grey', linestyle='--', linewidth=0.5)

# Cambiar el color de todos los bordes (spines) a gris
for spine in ax.spines.values():
    spine.set_color('#363636')

# Configurar los parámetros de las marcas y la cuadrícula
ax.tick_params(colors='0.2', grid_color='#515151')

# Ajustar los márgenes de la figura
plt.subplots_adjust(left=0.05, right=0.95, top=0.95, bottom=0.05)

#Ajustar el tamaño de la figura
plt.tight_layout()

# Guardar la imagen
plt.savefig('mapa.png')
plt.close()

#### PERFIL DEMOGRÁFICO ####

# Aplicar el filtro del municipio seleccionado y del año
datos_municipio_población = socioeconomico_quindio[(socioeconomico_quindio['Municipio'] == municipio_seleccionado)
                                                    & (socioeconomico_quindio['Año'] >= 2013)]

# definir las variables
poblacion_rural = dato_población_quindio['Población Rural'].values[0]
porcentaje_poblacion_rural = round(poblacion_rural / poblacion_total *100, 2)

poblacion_urbana = dato_población_quindio['Población Urbana'].values[0]
porcentaje_poblacion_urbana = round(poblacion_urbana / poblacion_total *100, 2)

#Gráfico


#Crear la figura y los ejes
plt.figure(figsize=(10, 5))
ax = plt.axes()

#Crear la gráfica de líneas para la población total
ax.plot(datos_municipio_población['Año'], 
         datos_municipio_población['Población Total'], 
         color='royalblue', 
         marker='o', 
         linestyle='-', 
         linewidth=2, 
         markersize=8,
         label='Población Total')

#Ajustar el ancho de las barras
bar_width = 0.48

#Calcular la posición de las barras
posicion_rural = datos_municipio_población['Año'] - bar_width / 2
posicion_urbana = datos_municipio_población['Año'] + bar_width / 2

#Crear las barras para las poblaciones rural y urbana
barras_rural = ax.bar(posicion_rural, 
                      datos_municipio_población['Población Rural'], 
                      bar_width, 
                      color='green', 
                      label='Población Rural')
barras_urbana = ax.bar(posicion_urbana, 
                       datos_municipio_población['Población Urbana'], 
                       bar_width, 
                       color='grey', 
                       label='Población Urbana')

# Agregar etiquetas a las barras
for barra in barras_rural:
    altura = barra.get_height()
    ax.text(barra.get_x() + barra.get_width() / 2, altura - 0.09 * altura, f'{(altura/1000).round(1)}', ha='center', va='bottom', fontsize=9.5, color='black')
for barra in barras_urbana:
    altura = barra.get_height()
    ax.text(barra.get_x() + barra.get_width() / 2, altura - 0.09 * altura, f'{(altura/1000).round(1)}', ha='center', va='bottom', fontsize=9.5, color='black')

# Agregar etiquetas a la línea
for x, y in zip(datos_municipio_población['Año'], datos_municipio_población['Población Total']):
    ax.annotate(f'{round(y/1000, 1)}', xy=(x, y), xytext=(0, 5), textcoords='offset points', ha='center', fontsize=10, color='royalblue')


#Agregar los detalles estéticos
plt.title(f'Evolución poblacional de {municipio_seleccionado}', fontsize=20, fontweight='bold', pad=40)

#Configurar el eje x
ax.set_xticks(datos_municipio_población['Año'])
ax.set_xlabel('Año', color='black', fontsize=12)
ax.tick_params(axis='x', length=0)
ax.set_xticklabels(datos_municipio_población['Año'], fontsize=11)

#Ocultar los valores y ticks del eje y
ax.set_yticklabels([])
ax.yaxis.set_ticks([])


# Eliminar los bordes
for spine in ax.spines.values():
    spine.set_visible(False)

#Agregar y configurar la leyenda
ax.legend(loc='upper center', fontsize=11, bbox_to_anchor=(0.5, 1.15), ncols=3)

#Añadir texto con aclaración sobre los valores de población
plt.text(0.21, 0.99, '*Población en miles de habitantes', transform=ax.transAxes, fontsize=9, color='#4B4B4B')

#Mostrar el gráfico
plt.tight_layout()

# Guardar la imagen
plt.savefig('grafico1.png')
plt.close()


#### PERFIL ECONÓMICO-PRODUCTIVO ####

#### CULTIVOS

#Filtrar los datos del municipio seleccionado y del año 2022
datos_municipio_cultivos = evas_quindio[(evas_quindio['Municipio'] == municipio_seleccionado)
                                        & (evas_quindio['Año'] == 2022)]

# Definir las variables
# Obtener la fila con el valor máximo de 'Área sembrada (ha)'
fila_cultivo_principal = datos_municipio_cultivos.loc[datos_municipio_cultivos['Área sembrada (ha)'].idxmax()]

# Extraer el valor de hectáreas y el nombre del cultivo
hectareas_cultivo_principal = fila_cultivo_principal['Área sembrada (ha)']
cultivo_principal = fila_cultivo_principal['Cultivo']

# Lista de cultivos
lista_de_cultivos = datos_municipio_cultivos['Cultivo'].unique()
lista_de_cultivos_str = ', '.join(lista_de_cultivos)

# Extensión de cultivos
suma_de_cultivos = datos_municipio_cultivos['Área sembrada (ha)'].sum()

#Seleccionar los 10 cultivos con mayor área sembrada
datos_municipio_cultivos = datos_municipio_cultivos.sort_values('Área sembrada (ha)', ascending=False).head(10)

# Gráfico

#Crear la figura y los ejes
plt.figure(figsize=(10, 5))
ax = plt.axes()

#Crear las barras horizontales
barras = ax.barh(datos_municipio_cultivos['Cultivo'],datos_municipio_cultivos['Área sembrada (ha)'], color='green', label='Área sembrada')

#Organizar de mayor a menor los cultivos de acuerdo con su área sembrada
plt.gca().invert_yaxis()

# Agregar etiquetas a las barras
for barra in barras:
    ancho = barra.get_width()
    ax.text(ancho, barra.get_y() + barra.get_height() / 2, f'{ancho:.1f}', ha='left', va='center', fontsize=10, color='black')

#Agregar los detalles estéticos
plt.title(f'Área sembrada (ha) por cultivo en {municipio_seleccionado}', fontsize=20, fontweight='bold')

#Ocultar los valores y ticks del eje y
ax.set_xticklabels([])
ax.xaxis.set_ticks([])
ax.tick_params(axis='y', length=0)

# Eliminar los bordes
for spine in ax.spines.values():
    spine.set_visible(False)

#Agregar y configurar la leyenda
plt.tight_layout()

# Guardar la imagen
plt.savefig('grafico2.png')
plt.close()


#### PIB MUNICIPAL

# Aplicar el filtro del municipio seleccionado y del año
datos_municipio_valor_agregado = socioeconomico_quindio[(socioeconomico_quindio['Municipio'] == municipio_seleccionado)
                                                        & (socioeconomico_quindio['Año'] >= 2013)]

# Definir las variables
datos_municipio_pib = socioeconomico_quindio[(socioeconomico_quindio['Municipio'] == municipio_seleccionado)
                                                        & (socioeconomico_quindio['Año'] == 2022)]

#Extraer el valor del PIB
PIB_primario = datos_municipio_pib['PIB Primario'].values[0]
PIB_secundario = datos_municipio_pib['PIB Secundario'].values[0]
PIB_terciario = datos_municipio_pib['PIB Terciario'].values[0]
PIB_total = datos_municipio_pib['PIB'].values[0]

#Porcentajes
porcentaje_PIB_primario = round(PIB_primario / PIB_total * 100, 2)
porcentaje_PIB_secundario = round(PIB_secundario / PIB_total * 100, 2)
porcentaje_PIB_terciario = round(PIB_terciario / PIB_total * 100, 2)

#Gráifico

# Crear la figura y los ejes
plt.figure(figsize=(10, 5))
ax = plt.axes()

# Crear las barras apiladas
barras_primario = ax.bar(datos_municipio_valor_agregado['Año'], 
                         datos_municipio_valor_agregado['PIB Primario'], 
                         color='green', 
                         label='PIB Primario')

barras_secundario = ax.bar(datos_municipio_valor_agregado['Año'], 
                           datos_municipio_valor_agregado['PIB Secundario'], 
                           bottom=datos_municipio_valor_agregado['PIB Primario'], 
                           color='grey', 
                           label='PIB Secundario')

barras_terciario = ax.bar(datos_municipio_valor_agregado['Año'], 
                          datos_municipio_valor_agregado['PIB Terciario'], 
                          bottom=datos_municipio_valor_agregado['PIB Primario'] + datos_municipio_valor_agregado['PIB Secundario'], 
                          color='royalblue', 
                          label='PIB Terciario')

# Etiquetas con porcentaje para cada barra
for barras, columna in zip([barras_primario, barras_secundario, barras_terciario], 
                           ['PIB Primario', 'PIB Secundario', 'PIB Terciario']):
    for barra in barras:
        altura = barra.get_height()
        # Usamos el valor central de cada barra para encontrar el año
        año = barra.get_x() + barra.get_width() / 2
        # Obtener el índice correspondiente al año
        año_idx = datos_municipio_valor_agregado.index[(datos_municipio_valor_agregado['Año'] == año)].tolist()[0]
        total = datos_municipio_valor_agregado.loc[año_idx, 'PIB']
        porcentaje = (altura / total) * 100
        # Ajustar la posición vertical de la etiqueta
        pos_y = barra.get_y() + barra.get_height() / 2
        ax.text(barra.get_x() + barra.get_width() / 2, 
                pos_y, 
                f'{porcentaje:.1f}%', 
                ha='center', 
                va='center', 
                fontsize=10, 
                color='white')
    
# Agregar los detalles estéticos
plt.title(f'PIB por sector en {municipio_seleccionado}', fontsize=20, fontweight='bold', pad=40)

# Configurar el eje x
ax.set_xticks(datos_municipio_valor_agregado['Año'])
ax.set_xlabel('Año', color='black', fontsize=12)
ax.tick_params(axis='x', length=0)
ax.set_xticklabels(datos_municipio_valor_agregado['Año'], fontsize=11)

# Ocultar los valores y ticks del eje y
ax.set_yticklabels([])
ax.yaxis.set_ticks([])

# Eliminar los bordes
for spine in ax.spines.values():
    spine.set_visible(False)

# Agregar y configurar la leyenda
ax.legend(loc='upper center', fontsize=11, bbox_to_anchor=(0.5, 1.15), ncol=3)

# Añadir texto con aclaración sobre los valores de población
plt.text(0.24, 0.99, '*Porcentaje PIB Municipal', transform=ax.transAxes, fontsize=9, color='#4B4B4B')

# Mostrar el gráfico
plt.tight_layout()

# Guardar la imagen
plt.savefig('grafico3.png')
plt.close()

#### INGRESOS PREDIALES

# Aplicar el filtro del municipio seleccionado y del año
datos_municipio_ingresos_predial = socioeconomico_quindio[(socioeconomico_quindio['Municipio'] == municipio_seleccionado)
                                        & (socioeconomico_quindio['Año'] == 2022)]

# definir las variables
ingresos_predial = datos_municipio_ingresos_predial['Ingresos Prediales'].values[0]
ingreso_total = datos_municipio_ingresos_predial['Ingresos Totales'].values[0]
porcentaje_ingreso_predial = round(ingresos_predial / ingreso_total * 100, 2)

#Gráfico

# Crear la figura y los ejes
plt.figure(figsize=(10, 5))
ax = plt.axes()

# Calcular los valores
ingresos_totales = datos_municipio_ingresos_predial['Ingresos Totales'].sum()
ingresos_prediales = datos_municipio_ingresos_predial['Ingresos Prediales'].sum()
otros_ingresos = ingresos_totales - ingresos_prediales

# Crear el gráfico de pastel
ax.pie([ingresos_prediales, otros_ingresos], labels=['Ingresos Prediales', 'Otros Ingresos'], autopct='%1.1f%%', colors=['#00CC66','#0066FF'], startangle=90)
ax.axis('equal')

# Agregar los detalles estéticos
plt.title(f'Proporción de Ingresos Prediales en {municipio_seleccionado}', fontsize=20, fontweight='bold')

# Mostrar el gráfico
plt.tight_layout()

# Guardar la imagen
plt.savefig('grafico4.png')
plt.close()

#### POBREZA MULTIDIMENSIONAL ####

#Aplicar el fitlro del municipio seleccionado
datos_municipio_pobreza = ipm_quindio[(ipm_quindio['Municipio'] == municipio_seleccionado)]

#Definir las variables IPM general
ipm_municipal = datos_municipio_pobreza['IPM'].values[0]
IPM_Rural = datos_municipio_pobreza['IPM Rural'].values[0]
IPM_Urbano = datos_municipio_pobreza['IPM Urbano'].values[0]

#Gráfico

# Crear la figura y los ejes
plt.figure(figsize=(10, 5))
ax = plt.axes()

# Preparar los datos para el gráfico
categorias = ['IPM Municipal', 'IPM Urbano', 'IPM Rural']
valores = [datos_municipio_pobreza['IPM'].values[0], datos_municipio_pobreza['IPM Urbano'].values[0], datos_municipio_pobreza['IPM Rural'].values[0]]

# Crear las barras horizontales
ax.barh(categorias, valores, color=['#0066FF','#808080','#00CC99'])

# Agregar el título
plt.title(f'Índice de Pobreza Multidimensional en {municipio_seleccionado}', fontsize=20, fontweight='bold')

# Agregar etiquetas a las barras
for i, v in enumerate(valores):
    ax.text(v + 0.01, i, f'{v:.2f}', color='black', va='center', fontsize=10)

#Ocultar los valores y ticks del eje y
ax.set_xticklabels([])
ax.xaxis.set_ticks([])
ax.tick_params(axis='y', length=0)

# Eliminar los bordes
for spine in ax.spines.values():
    spine.set_visible(False)

# Mostrar el gráfico
plt.tight_layout()

# Guardar la imagen
plt.savefig('grafico5.png')
plt.close()

#Definir las variables mayores IPM
trabajo_informal_municipio = datos_municipio_pobreza['Trabajo Informal'].values[0]
trabajo_informal_rural = datos_municipio_pobreza['Trabajo Informal Rural'].values[0]
trabajo_informal_urbano = datos_municipio_pobreza['Trabajo Informal Urbano'].values[0]


##### DOCUMENTO DE WORD ####

doc = Document()

# Ajustar los márgenes del documento
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

# Función para agregar párrafos justificados con tamaño de letra 12
def add_paragraph(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph.runs:
        run.font.size = Pt(12)

# Agregar contenido al documento
doc.add_heading(f'Caracterización Territorial del Municipio {municipio_seleccionado}, Quindío', 0)
doc.add_heading('Introducción', level=1)
add_paragraph(doc, f'El municipio de {municipio_seleccionado} se encuentra ubica en el {ubicaciones[municipio_seleccionado]} del Departamento tiene una extensión de {extension} km2 y una población de {poblacion_total} en el año 2022 (DANE,2022). A continuación, se muestra su ubicación en el Departamento:')
doc.add_picture('mapa.png', width=Inches(6.5))
doc.add_heading('Perfil Demográfico', level=1)
add_paragraph(doc, f'De acuerdo con el DANE, {municipio_seleccionado} al año 2022 tiene una población de {poblacion_total} que se distribuye en {poblacion_rural} habitantes rurales, equivalente al {porcentaje_poblacion_rural}% del total y una población urbana de {poblacion_urbana} habitantes el {porcentaje_poblacion_urbana}%. La gráfica 1 muestra la evolución de la población de {municipio_seleccionado} entre los años 2013 a 2022:')
doc.add_picture('grafico1.png', width=Inches(6.5))
doc.add_heading('Perfil Económico-Productivo', level=1)
doc.add_heading('Cultivos', level=2)
add_paragraph(doc, 'La economía del departamento del Quindío se caracteriza por su orientación a los sectores primario y terciario de la economía, específicamente al desarrollo de actividades agropecuarias y la prestación de servicios enfocados en el turismo y sus derivados.')
add_paragraph(doc, f'De acuerdo con las Evaluaciones Agropecuarias Municipales, EVAS, del Ministerio de Agricultura y Desarrollo Rural y la UPRA, los cultivos presentes en {municipio_seleccionado}, durante el año 2022, son {lista_de_cultivos_str} con una extensión conjunta de {suma_de_cultivos} hectáreas, de estos el principal es {cultivo_principal} con una extensión de {hectareas_cultivo_principal} hectáreas. La gráfica 2 presenta principales cultivos al año 2022, de acuerdo con su área sembrada:')
doc.add_picture('grafico2.png', width=Inches(6.5))
doc.add_heading('PIB Municipal', level=2)
add_paragraph(doc, f'Las Cuentas Nacionales del DANE estiman el Producto Interno Bruto Municipal a partir del valor agregado generado por las actividades en los tres sectores de la economía. En el caso de {municipio_seleccionado}, durante el año 2022, el sector primario representó el {porcentaje_PIB_primario}%, el sector secundario tuvo un peso del {porcentaje_PIB_secundario}% y el sector terciario generó el {porcentaje_PIB_terciario}% del total. La gráfica 3 muestra la evolución del PIB entre los años 2013 a 2022, con el porcentaje de cada uno de los tres sectores:')
doc.add_picture('grafico3.png', width=Inches(6.5))
doc.add_heading('Ingresos Prediales', level=2)
add_paragraph(doc, f'Los Ingresos Municipales dependen de diversos factores, transferencias nacionales, actividades económicas del municipio, regalías, entre otros, sin embargo, para la mayoría de los municipios de sexta categoría (la predominante en Quindío), el ingreso por impuesto predial representa la mayor parte de sus Ingresos Corrientes de Libre Destinación, es decir, aquellos de los que pueden disponer para realizar las inversiones que requiere su municipio y de los cuales tiene libertad para decidir su destinación. Para el caso {municipio_seleccionado}, los ingresos por predial representan el {porcentaje_ingreso_predial}% de los ingresos totales como se observa en el gráfico 4:')
doc.add_picture('grafico4.png', width=Inches(6.5))
doc.add_heading('Pobreza Multidimensional', level=2)
add_paragraph(doc, 'El DANE mide el Indicador de Pobreza Multidimensional, IPM, a partir de las privaciones que tienen los hogares. A partir de los datos del Censo 2018, el DANE, realizó mediciones precisas de las diferentes dimensiones que componen este indicador brindando un panorama muy detallado de las atenciones que necesitan los municipios. El indicador se mide de 1 a 100, representando el porcentaje de las personas que se encuentran en condición de privación.')
add_paragraph(doc, f'{municipio_seleccionado} tiene un IPM de {ipm_municipal}, sin embargo, al territorializarlo se identifica que la situación es más difícil en el sector rural con un IPM de {IPM_Rural} en comparación con el IPM urbano que es de {IPM_Urbano}. La gráfica 5 muestra los niveles de IPM, por sector en {municipio_seleccionado}:')
doc.add_picture('grafico5.png', width=Inches(6.5))
add_paragraph(doc, f'La principal problemática que impacta a una mayor proporción de la población en términos de pobreza multidimensional del municipio es el trabajo informal, a nivel municipal es de {trabajo_informal_municipio}, en el sector rural es de {trabajo_informal_rural} y en el sector urbano es de {trabajo_informal_urbano}.')
doc.save(f'{municipio_seleccionado}.docx')